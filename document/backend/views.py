from django.shortcuts import render
import warnings
warnings.simplefilter(action='ignore')
from django.shortcuts import redirect
from django.http import HttpResponse
import docx
from docx import Document
from docx.shared import Mm,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import os
from docx.shared import Mm,Pt,Inches
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx import *
from docx.text.paragraph import Paragraph
from docx.text.paragraph import Run
import xml.etree.ElementTree as ET
from docx.document import Document as doctwo
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt
from docxcompose.composer import Composer
from docx import Document as Document_compose
import pandas as pd
from xml.etree import ElementTree
from io import StringIO
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from io import BytesIO
import csv
import base64
from docx.enum.table import WD_ALIGN_VERTICAL
import pickle
import spacy



def home(request):
    return render(request,"home.html")

def process(request):
    if request.method == "POST":
        file = request.FILES["myFile"]
        # existing_doc = Document(file)
        document = Document(file)
        # e3s_doc = Document()
        # section = e3s_doc.sections[0]
        # section.start_type
        # section.page_width = Mm(170)
        # section.page_height = Mm(250)
        # section.left_margin = Mm(20)
        # section.right_margin = Mm(20)
        # section.top_margin = Mm(24)
        # section.bottom_margin = Mm(16)
        # title = False
        # abs = r'^Abstract\b.*'
        # for element in existing_doc:
        #     if(isinstance(element, docx.text.paragraph.Paragraph)):
        #         if title == False:
        #             e3s_paragraph = e3s_doc.add_paragraph((element.text).title())
        #             e3s_paragraph.paragraph_format.space_before = Mm(22)
        #             e3s_paragraph.paragraph_format.space_before = Mm(6) 
        #             for run in e3s_paragraph.runs:
        #                 run.font.name = 'Arial'
        #                 run.font.size = Pt(16)
        #                 run.bold = True
        #             title = True
        #             e3s_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        #             continue
        #         # Create a new paragraph in the E3S-formatted document
        #         e3s_paragraph = e3s_doc.add_paragraph(element.text)

        #         s = e3s_paragraph.text
        #         for run in e3s_paragraph.runs:
        #             run.font.name = 'Times New Roman'

        #         if re.search(abs,s):
        #             split_index = s.find("Abstract") + len("Abstract")
        #             first_half_text = s[:split_index]
        #             second_half_text = s[split_index:]
        #             e3s_paragraph.clear()   
        #             run1 = e3s_paragraph.add_run(first_half_text)
        #             run1.bold = True
        #             run1.font.name = 'Arial'
        #             run2 = e3s_paragraph.add_run(second_half_text)
        #             run2.font.name = 'Times New Roman'
        #         e3s_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        #     elif isinstance(element, docx.text.run.InlineShape):
        #         print("picture")
        #         image_blob = element.blob
        #         e3s_doc.add_picture(image_blob)
        #     elif isinstance(element, docx.table.Table):
        #         print("This is a table.")
        # e3s_doc.save('formatted_document.docx')
        nlp = spacy.load("en_core_web_sm")
        with open(r"C:\Users\nysha\gradient_boost_model.pkl", 'rb') as f:
            loaded_model = pickle.load(f)
        def ml_check(para):
            word_c = len(para.split(' '))
            try:
                if int(loaded_model.predict([[len(para), word_c]]))==1:
                    return True
                else:
                    return False
            except:
                return False
            
        def contains_person_name(sentence):
            doc = nlp(sentence)
            for ent in doc.ents:
                if ent.label_ == "PERSON":
                    return True
            return False 

        def iter_block_items(parent):
            """
            Yield each paragraph and table child within *parent*, in document order.
            Each returned value is an instance of either Table or Paragraph. *parent*
            would most commonly be a reference to a main Document object, but
            also works for a _Cell object, which itself can contain paragraphs and tables.
            """
            if isinstance(parent, doctwo):
                parent_elm = parent.element.body
            elif isinstance(parent, _Cell):
                parent_elm = parent._tc
            else:
                raise ValueError("something's not right")

            for child in parent_elm.iterchildren():
                if isinstance(child, CT_P):
                    yield Paragraph(child, parent)
                elif isinstance(child, CT_Tbl):
                    yield Table(child, parent)

            
        #This function extracts the table from the document object as a dataframe
        def read_docx_tables(tab_id=None, **kwargs):
            """
            parse table(s) from a Word Document (.docx) into Pandas DataFrame(s)

            Parameters:
                filename:   file name of a Word Document

                tab_id:     parse a single table with the index: [tab_id] (counting from 0).
                            When [None] - return a list of DataFrames (parse all tables)

                kwargs:     arguments to pass to `pd.read_csv()` function

            Return: a single DataFrame if tab_id != None or a list of DataFrames otherwise
            """
            def read_docx_tab(tab, **kwargs):
                vf = io.StringIO()
                writer = csv.writer(vf)
                for row in tab.rows:
                    writer.writerow(cell.text for cell in row.cells)
                vf.seek(0)
                return pd.read_csv(vf, **kwargs)

        #    doc = Document(filename)
            if tab_id is None:
                return [read_docx_tab(tab, **kwargs) for tab in document.tables]
            else:
                try:
                    return read_docx_tab(document.tables[tab_id], **kwargs)
                except IndexError:
                    print('Error: specified [tab_id]: {}  does not exist.'.format(tab_id))
                    raise



        #The combined_df dataframe will store all the content in document order including images, tables and paragraphs.
        #If the content is an image or a table, it has to be referenced from image_df for images and table_list for tables using the corresponding image or table id that is stored in combined_df
        #And if the content is paragraph, the paragraph text will be stored in combined_df
        combined_df = pd.DataFrame(columns=['para_text','table_id','style'])
        table_mod = pd.DataFrame(columns=['string_value','table_id'])

        #The image_df will consist of base64 encoded image data of all the images in the document
        image_df = pd.DataFrame(columns=['image_index','image_rID','image_filename','image_base64_string'])

        #The table_list is a list consisting of all the tables in the document
        table_list=[]
        xml_list=[]
        title = False
        i=0
        imagecounter = 0
        table_counter = 0
        abs = r'^ABSTRACT\b.*'
        gmail_pattern = pattern = r'\@\w+\.com' 

        blockxmlstring = ''
        for block in iter_block_items(document):
            if 'text' in str(block):
                isappend = False
                runboldtext = ''
                for run in block.runs:                        
                    if run.bold:
                        runboldtext = runboldtext + run.text
                
                style = str(block.style.name)
                appendtxt = str(block.text)
                appendtxt = appendtxt.replace("\n","")
                appendtxt = appendtxt.replace("\r","")
                tabid = 'Novalue'
                paragraph_split = appendtxt.lower().split()                
                isappend = True
                for run in block.runs:
                    xmlstr = str(run.element.xml)
                    my_namespaces = dict([node for _, node in ElementTree.iterparse(StringIO(xmlstr), events=['start-ns'])])
                    root = ET.fromstring(xmlstr) 
                    #Check if pic is there in the xml of the element. If yes, then extract the image data
                    if 'pic:pic' in xmlstr:
                        xml_list.append(xmlstr)
                        for pic in root.findall('.//pic:pic', my_namespaces):
                            cNvPr_elem = pic.find("pic:nvPicPr/pic:cNvPr", my_namespaces)
                            name_attr = cNvPr_elem.get("name")
                            blip_elem = pic.find("pic:blipFill/a:blip", my_namespaces)
                            embed_attr = blip_elem.get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                            isappend = True
                            appendtxt = str('Document_Imagefile/' + name_attr + '/' + embed_attr + '/' + str(imagecounter))
                            document_part = document.part
                            image_part = document_part.related_parts[embed_attr]
                            image_base64 = base64.b64encode(image_part._blob)
                            image_base64 = image_base64.decode()                            
                            dftemp = pd.DataFrame({'image_index':[imagecounter],'image_rID':[embed_attr],'image_filename':[name_attr],'image_base64_string':[image_base64]})
                            image_df = image_df.append(dftemp,sort=False)
                            style = 'Novalue'
                        imagecounter = imagecounter + 1
                    
            elif 'table' in str(block):
                isappend = True
                style = 'Novalue'
                appendtxt = str(block)
                tabid = i
                dfs = read_docx_tables(tab_id=i)
                dftemp = pd.DataFrame({'para_text':[appendtxt],'table_id':[i],'style':[style]})
                table_mod = table_mod.append(dftemp,sort=False)
                table_list.append(dfs)
                i=i+1
            if isappend:
                    dftemp = pd.DataFrame({'para_text':[appendtxt],'table_id':[tabid],'style':[style]})
                    combined_df=combined_df.append(dftemp,sort=False)
                    
        combined_df = combined_df.reset_index(drop=True)
        image_df = image_df.reset_index(drop=True)

        print(combined_df)
        # print(combined_df[combined_df["style"] == "Novalue"])
        print(combined_df["style"])
        # print(combined_df[combined_df["style"] == "Affiliation"])
        # print(image_df.head())

        # print(combined_df)

        # print(table_list[0].head())



        def check(text):
            l = list(text.split(" "))
            if len(l)>9:
                return False
            return True


        # print(combined_df[combined_df['para_text'].str.contains("reference", case=False)])

        target_doc = Document()
        section = target_doc.sections[0]
        section.start_type
        section.page_width = Mm(170)
        section.page_height = Mm(250)
        section.left_margin = Mm(20)
        section.right_margin = Mm(20)
        section.top_margin = Mm(24)
        section.bottom_margin = Mm(16)
        image_idx = 0
        rows = combined_df.shape[0]

        for row in range(rows):
            # Identifying a paragraph to the document
            if combined_df.iloc[row]["style"] != "Novalue":
                # target_doc.add_paragraph(combined_df.iloc[row]["para_text"])

                combined_df.iloc[row]["para_text"] = combined_df.iloc[row]["para_text"].strip()
                if title == False:
                    para = target_doc.add_paragraph((combined_df.iloc[row]["para_text"]).title())
                    para.paragraph_format.space_before = Mm(22)
                    para.paragraph_format.space_after = Mm(6) 
                    for run in para.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(16)
                        run.bold = True
                    title = True
                elif re.search(abs,combined_df.iloc[row]["para_text"],re.IGNORECASE):
                    # print("Identified Abstract")
                    split_index = combined_df.iloc[row]["para_text"].find("Abstract") + len("Abstract")
                    first_half_text = combined_df.iloc[row]["para_text"][:split_index+2]
                    second_half_text = combined_df.iloc[row]["para_text"][split_index+2:]  
                    para = target_doc.add_paragraph()
                    run1 = para.add_run(first_half_text)
                    run1.font.name = 'Arial'
                    run1.font.size = Pt(10)
                    run1.bold = True
                    run2 = para.add_run(second_half_text)
                    run2.font.name = 'Times New Roman'
                    run2.font.size = Pt(9)
                elif ((check(combined_df.iloc[row]["para_text"])==True or ml_check(combined_df.iloc[row]["para_text"])==True) and (contains_person_name(combined_df.iloc[row]["para_text"])==True) or (re.search(gmail_pattern,combined_df.iloc[row]["para_text"],re.IGNORECASE))):
                    para = target_doc.add_paragraph((combined_df.iloc[row]["para_text"]))
                    for run in para.runs:
                        run.font.name =  'Times New Roman'
                        run.font.size = Pt(9)
                        run.italic = True
                elif  combined_df.iloc[row]["style"] == "Title":
                    para = target_doc.add_paragraph((combined_df.iloc[row]["para_text"]).title())
                    para.paragraph_format.space_before = Mm(22)
                    para.paragraph_format.space_after = Mm(6) 
                    for run in para.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(16)
                        run.bold = True
                    title = True
                
                elif  combined_df.iloc[row]["style"] == "Section":
                    para = target_doc.add_paragraph(combined_df.iloc[row]["para_text"])
                    for run in para.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(12)
                        run.bold = True
                elif combined_df.iloc[row]["style"] == "Affiliation":
                    para = target_doc.add_paragraph(combined_df.iloc[row]["para_text"])
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(9)
                elif combined_df.iloc[row]["style"] == "Author Last Name":
                    para = target_doc.add_paragraph(combined_df.iloc[row]["para_text"])
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(9)
                        run.italic = True
                elif combined_df.iloc[row]["style"] == "Style Figure Caption" or combined_df.iloc[row]["style"] == "Table Caption" :
                    para = target_doc.add_paragraph(combined_df.iloc[row]["para_text"])
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(9)
                elif  combined_df.iloc[row]["style"] == "Subsection" or combined_df.iloc[row]["style"] == "Subsubsection" or combined_df.iloc[row]["style"] == "Body Text" or (check(combined_df.iloc[row]["para_text"])==True or ml_check(combined_df.iloc[row]["para_text"]))and not combined_df.iloc[row]["style"].startswith("R"):
                    para = target_doc.add_paragraph(combined_df.iloc[row]["para_text"])
                    for run in para.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
                        run.bold = True
                else:
                    para = target_doc.add_paragraph(combined_df.iloc[row]["para_text"])
                    for run in para.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                    
                para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            # image is identified
            elif combined_df.iloc[row]["table_id"] == "Novalue":
                str_img = image_df.iloc[image_idx]["image_base64_string"]
                image_idx +=1  
                image_bytes = base64.b64decode(str_img)
                image_stream = BytesIO(image_bytes)
                target_doc.add_picture(image_stream, width=Inches(5.0))
        #     # Table is identified
            elif combined_df.iloc[row]["table_id"] != "Novalue":
                temp_df =  table_list[table_counter]
                table_counter+=1  
                table = target_doc.add_table(rows=1, cols=len(temp_df.columns))
                hdr_cells = table.rows[0].cells
                for i, col_name in enumerate(temp_df.columns):
                    hdr_cells[i].text = col_name

                # Add data rows
                for _, row in temp_df.iterrows():
                    new_row = table.add_row().cells
                    for i, val in enumerate(row):
                        new_row[i].text = str(val)


                for row in table.rows:
                    for cell in row.cells:
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(10)  
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                table.style = 'Table Grid'
        # table_obj = docx.table.Table(combined_df.iloc[row]["para_text"])
        # target_doc.add_table(table_obj)

        target_doc.save('target.docx')
        return redirect('/download')
    else:
        return HttpResponse("<h1> Some Technical issue </h1>")
    
    
    
def download(request):
    file_path = os.path.join(os.getcwd(),"target.docx")
    if os.path.exists(file_path):
        with open(file_path, 'rb') as file:
            response = HttpResponse(file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'inline; filename="doc.docx"'
            return response
    else:
        return render(request, 'error_page.html', {'error_message': 'File not found'})