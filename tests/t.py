from docx import Document
from docx.shared import Mm,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
# Load the existing Word document
# existing_doc = Document(r'C:\Users\nysha\Downloads\manuscript_A08.docx')
existing_doc = Document(r'C:\Users\nysha\Downloads\manuscript_A08.docx')

# Create a new Word document for the E3S-formatted content
e3s_doc = Document()



# Set page margins for the E3S-style document (adjust as needed)
section = e3s_doc.sections[0]
section.start_type
section.page_width = Mm(170)
section.page_height = Mm(250)
section.left_margin = Mm(20)
section.right_margin = Mm(20)
section.top_margin = Mm(24)
section.bottom_margin = Mm(16)

title = False

# Iterate through paragraphs in the existing document and apply E3S formatting
abs = r'^Abstract\b.*'
for paragraph in existing_doc.paragraphs:
    if title == False:
        e3s_paragraph = e3s_doc.add_paragraph((paragraph.text).title())
        e3s_paragraph.paragraph_format.space_before = Mm(22)
        e3s_paragraph.paragraph_format.space_before = Mm(6) 
        for run in e3s_paragraph.runs:
            run.font.name = 'Arial'
            run.font.size = Pt(16)
            run.bold = True
        title = True
        e3s_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        continue
    # Create a new paragraph in the E3S-formatted document
    e3s_paragraph = e3s_doc.add_paragraph(paragraph.text)

    s = e3s_paragraph.text
    for run in e3s_paragraph.runs:
        run.font.name = 'Times New Roman'

    if re.search(abs,s):
        split_index = s.find("Abstract") + len("Abstract")
        first_half_text = s[:split_index]
        second_half_text = s[split_index:]
        e3s_paragraph.clear()   
        run1 = e3s_paragraph.add_run(first_half_text)
        run1.bold = True
        run1.font.name = 'Arial'
        run2 = e3s_paragraph.add_run(second_half_text)

        

    # Set paragraph alignment to left
    e3s_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    

# Save the E3S-formatted document
e3s_doc.save('doc.docx')