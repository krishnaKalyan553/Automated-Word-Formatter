# from docx import Document

# # Open the Word document
# doc = Document(r'C:\Users\nysha\Downloads\manuscript_A08.docx')

# # Initialize a list to store extracted paragraphs
# paragraphs = []

# # Iterate through paragraphs and add them to the list
# for paragraph in doc.paragraphs:
#     paragraphs.append(paragraph.text)

# # Print or process the extracted paragraphs as needed
# print(paragraphs[1])



# from docx import Document
# from docx.shared import Pt
# from docx.enum.text import WD_ALIGN_PARAGRAPH

# # Load the existing Word document
# existing_doc = Document(r'C:\Users\nysha\Downloads\manuscript_A08.docx')

# # Create a new Word document for the E3S-formatted content
# e3s_doc = Document()

# # Set page margins for the E3S-style document (adjust as needed)
# section = e3s_doc.sections[0]
# section.left_margin = Pt(2.54)
# section.right_margin = Pt(2.54)
# section.top_margin = Pt(2.54)
# section.bottom_margin = Pt(2.54)


# # Iterate through paragraphs in the existing document and apply E3S formatting
# for paragraph in existing_doc.paragraphs:
#     # Create a new paragraph in the E3S-formatted document
#     e3s_paragraph = e3s_doc.add_paragraph(paragraph.text)
#     # Check for specific headings and apply formatting (adjust as needed)
#     if "Abstract" in paragraph.text:
#         e3s_paragraph.style = 'Heading1'  # Apply the style for abstract headings
#     elif "Introduction" in paragraph.text:
#         e3s_paragraph.style = 'Heading1'  # Apply the style for introduction headings
#     # Add more conditions for other headings as needed

#     # Set paragraph alignment to left
#     e3s_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

# # Save the E3S-formatted document
# e3s_doc.save('e3s_formatted_document.docx')



from docx import Document

# Load the existing Word document
existing_doc = Document(r'C:\Users\nysha\Downloads\manuscript_A08.docx')

# Create a new Word document for the cloned content
cloned_doc = Document()

# Iterate through sections in the existing document and copy them to the new document
l = []
for section in existing_doc.sections:
    cloned_section = cloned_doc.add_section()
    cloned_section.start_type
    cloned_section.orientation = section.orientation
    cloned_section.page_width = section.page_width
    cloned_section.page_height = section.page_height
    cloned_section.left_margin = section.left_margin
    cloned_section.right_margin = section.right_margin
    cloned_section.top_margin = section.top_margin
    cloned_section.bottom_margin = section.bottom_margin
    cloned_section.header_distance = section.header_distance
    cloned_section.footer_distance = section.footer_distance
    cloned_section.gutter = section.gutter

Iterate through elements in the existing document and copy them to the new document
for element in existing_doc.element.body:
    if element.tag.endswith('sectPr'):
        continue  # Skip section properties
    cloned_doc.element.body.append(element)

# Save the cloned document
cloned_doc.save('cloned_document.docx')
