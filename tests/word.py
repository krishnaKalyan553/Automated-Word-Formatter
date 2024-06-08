# import docx
# # Create a new Word document

# # print(dir(docx.enum.text))

# doc = docx.Document()

# #  IEEE-style title
# t = input("Enter the Title of Your IEEE Paper \n")
# title = doc.add_heading(t , level=1)
# title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

# # Authors
# authors = doc.add_paragraph('Author(s): Author Name(s)')
# authors.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

# abs = input("Enter the Abstract of Your IEEE Paper \n")
# # Abstract
# doc.add_heading('Abstract', level=1)
# abstract = doc.add_paragraph(abs)

# #  IEEE-style References section
# doc.add_heading('References', level=1)

# # Add IEEE-style references
# references = [
#     "Author1, A., Author2, B., & Author3, C. (Year). Title of the Paper. Journal, Volume(Issue), Page-Page.",
#     "Another Author, D., & More Authors, E. (Year). Another Title. Conference Name, Page-Page."
# ]

# for ref in references:
#     doc.add_paragraph(ref, style='ListBullet')

# #  Save the document
# doc.save('p.docx')



# from docx import Document
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.shared import Pt

# # Create a new Word document
# doc = Document()

# # Title
# title = doc.add_heading('IEEE-Style Document Example', level=1)
# title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# # Authors
# authors = doc.add_paragraph()
# authors.add_run('Author 1').bold = True
# authors.add_run(', Author 2, Author 3')
# authors.alignment = WD_ALIGN_PARAGRAPH.CENTER

# # Abstract
# abstract = doc.add_paragraph()
# abstract.add_run('Abstract:').bold = True
# abstract.add_run(' This is the abstract of your paper. It should provide a brief summary of your research.').italic = True

# # Section 1: Introduction
# section1 = doc.add_heading('I. INTRODUCTION', level=2)
# section1.alignment = WD_ALIGN_PARAGRAPH.LEFT

# intro_text = doc.add_paragraph()
# intro_text.add_run('This is the introduction section of your paper.').bold = True

# # Section 2: Methodology
# section2 = doc.add_heading('II. METHODOLOGY', level=2)
# section2.alignment = WD_ALIGN_PARAGRAPH.LEFT

# method_text = doc.add_paragraph()
# method_text.add_run('This is the methodology section of your paper.').bold = True

# # Section 3: Results
# section3 = doc.add_heading('III. RESULTS', level=2)
# section3.alignment = WD_ALIGN_PARAGRAPH.LEFT

# results_text = doc.add_paragraph()
# results_text.add_run('This is the results section of your paper.').bold = True

# # References
# references = doc.add_heading('REFERENCES', level=2)
# references.alignment = WD_ALIGN_PARAGRAPH.LEFT

# # Add a sample reference
# reference1 = doc.add_paragraph()
# reference1.add_run('[1] Author A, Author B, "Title of Paper," Journal Name, vol. 1, no. 1, pp. 1-10, Year.').font.size = Pt(10)

# # Save the document
# doc.save('IEEE_style_document.docx')





from docx import *

# # Open the Word document
# doc = Document(r'C:\Users\nysha\OneDrive\MINI PROJECT.docx')

# # Initialize variables for title, authors, and abstract
# title = ""
# authors = ""
# abstract = ""

# # Iterate through paragraphs and identify specific information
# for paragraph in doc.paragraphs:
#     if "Title:" in paragraph.text:
#         title = paragraph.text.replace("Title:", "").strip()
#     elif "Authors:" in paragraph.text:
#         authors = paragraph.text.replace("Authors:", "").strip()
#     elif "Abstract:" in paragraph.text:
#         abstract = paragraph.text.replace("Abstract:", "").strip()

# # Create a new Word document
ieee_doc = Document()

ieee_doc.add
# # IEEE-Style Title
# ieee_title = ieee_doc.add_heading(title, level=1)
# ieee_title.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

# # IEEE-Style Authors
# ieee_authors = ieee_doc.add_paragraph(authors)
# ieee_authors.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.CENTER

# # IEEE-Style Abstract
# ieee_abstract = ieee_doc.add_paragraph('Abstract:', style='List Bullet')
# ieee_abstract.paragraph_format.left_indent = docx.shared.Inches(0.5)
# ieee_abstract.add_run(abstract).italic = True

# Save the IEEE-formatted document
ieee_doc.save('ieee_formatted_document.docx')
